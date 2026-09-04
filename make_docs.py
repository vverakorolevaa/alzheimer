"""
Генерация литературного обзора проекта (docs/литературный_обзор.docx).

Запуск: python3 make_docs.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS = "docs"
os.makedirs(DOCS, exist_ok=True)


def _set_font(run, size=12, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def _para(doc, text="", bold_prefix=None, size=12):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        _set_font(r, size=size, bold=True, color=(0x2C, 0x3E, 0x50))
    if text:
        r = p.add_run(text)
        _set_font(r, size=size)
    return p


def _bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_font(r, size=size)
    return p


def _table(doc, headers, rows, size=11):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            _set_font(run, size=size, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
            for run in table.rows[r_idx + 1].cells[c_idx].paragraphs[0].runs:
                _set_font(run, size=size)
    return table


def make_litreview():
    doc = Document()
    doc.core_properties.title = "Литературный обзор"

    _heading(doc, "Литературный обзор", level=1)
    _para(doc, "Стадийная диагностика болезни Альцгеймера по экспрессии генов крови "
               "с помощью методов машинного обучения", size=12)
    doc.add_paragraph()

    # 1
    _heading(doc, "1. Болезнь Альцгеймера: масштаб проблемы и патогенез", level=2)
    _para(doc, "Болезнь Альцгеймера — наиболее распространённая форма деменции: по данным ВОЗ "
               "(2023) деменцией страдают более 55 млн человек, на БА приходится 60–70% случаев. "
               "Согласно амилоидной гипотезе (Selkoe, Hardy, 2016), ключевое звено патогенеза — "
               "накопление бляшек бета-амилоида вне нейронов и клубков тау-белка внутри них. "
               "Braak и Braak (1991) показали, что патология закономерно распространяется от "
               "энторинальной коры и гиппокампа к новой коре; современное нейропатологическое "
               "описание дано DeTure и Dickson (2019).")
    _para(doc, "Молекулярные изменения начинаются за 10–20 лет до симптомов. На этом основана "
               "биомаркерная классификация NIA-AA (Jack et al., 2018): болезнь — непрерывный "
               "континуум от асимптоматической стадии через лёгкие когнитивные нарушения (MCI) "
               "к деменции. Отсюда центральная задача: выявлять болезнь в «окне» MCI.")

    # 2
    _heading(doc, "2. Клинические стадии и нейропсихологические шкалы", level=2)
    _para(doc, "Концепцию MCI как переходной стадии систематизировал Petersen с соавт. (2001); "
               "ежегодная конверсия MCI в деменцию достигает 10–15%. Стадию деменции "
               "устанавливают шкалы MMSE (Folstein et al., 1975; 0–30 баллов) и CDR "
               "(Hughes et al., 1982; Morris, 1993). Результаты тестов зависят от образования "
               "и когнитивного резерва (Stern, 2012), поэтому объективные молекулярные "
               "биомаркеры необходимы как дополнение.")

    # 3
    _heading(doc, "3. Почему кровь: транскриптомика периферической крови при БА", level=2)
    _para(doc, "Существующие молекулярные биомаркеры БА требуют дорогих или инвазивных "
               "процедур: ПЭТ-сканирование, люмбальная пункция. Кровь же доступна у любого "
               "пациента при рутинном заборе, что делает её идеальным материалом для скрининга. "
               "Хотя нервная ткань отделена гематоэнцефалическим барьером, при БА в крови "
               "воспроизводимо меняется экспрессия сотен генов — отражение системного "
               "воспаления, окислительного стресса и энергетического дефицита.")
    _para(doc, bold_prefix="Ключевая работа: ",
          text="Lunnon с соавт. (2013) на когорте AddNeuroMed построили первый классификатор "
               "болезни Альцгеймера по экспрессии генов цельной крови и показали, что он "
               "помечает часть пациентов с MCI как «AD-подобных» — то есть кровь чувствительна "
               "к ранней стадии. Именно эта когорта (GSE63060, платформа Illumina HumanHT-12 "
               "V3.0) используется в настоящем проекте.")
    _para(doc, "Lee и Lee (2020) обучили ансамбль моделей на датасетах GSE63060/GSE63061 и "
               "подтвердили предсказуемость БА по крови (AUC порядка 0.8), отметив, что "
               "наибольшую устойчивость дают гены энергетического обмена и рибосомальные гены "
               "— та же биология, что воспроизводит и наша панель.")

    # 4
    _heading(doc, "4. Терапевтическое окно: почему ранняя стадия решает всё", level=2)
    _para(doc, "Моноклональные антитела к бета-амилоиду леканемаб (van Dyck et al., 2023) и "
               "донанемаб (Sims et al., 2023) замедляют снижение когнитивных функций на 27–35%, "
               "но работают только на ранних стадиях (MCI и лёгкая деменция). Таким образом, "
               "ценность диагностического инструмента определяется тем, насколько рано и точно "
               "он разделяет норма / MCI / деменция.")

    # 5
    _heading(doc, "5. Машинное обучение в диагностике БА", level=2)
    _para(doc, "Большинство работ по МО-диагностике БА опирается на нейровизуализацию "
               "(МРТ/ПЭТ) и закрытые данные ADNI: Zhang et al. (2011) — мультимодальный SVM "
               "(AUC 0.93); Liu et al. (2018) — каскадные CNN на МРТ+ПЭТ; Spasov et al. (2019) — "
               "параметр-эффективная 3D-CNN для прогноза конверсии MCI→AD.")
    _table(doc,
           ["Работа", "Данные", "Метод", "Ограничение"],
           [
               ["Zhang et al., 2011", "ADNI: МРТ+ПЭТ+ликвор", "SVM с мультимодальным ядром",
                "Дорогие/инвазивные модальности; закрытый доступ"],
               ["Liu et al., 2018", "ADNI: МРТ+ПЭТ", "Каскадные CNN",
                "Нужны полные снимки; слабая интерпретируемость"],
               ["Spasov et al., 2019", "ADNI: МРТ+ПЭТ", "3D-CNN", "ADNI; объёмные снимки"],
               ["Lunnon et al., 2013", "AddNeuroMed: кровь", "Селекция генов + классификатор",
                "2 класса (AD/контроль); без веб-инструмента"],
               ["Lee, Lee, 2020", "GSE63060/63061: кровь", "Ансамбль МО",
                "2 класса; нет персонального объяснения"],
               ["Настоящая работа", "GSE63060: кровь (открытые)", "Логрегрессия по мини-панели + SHAP",
                "Одна когорта; MCI трудный класс"],
           ])
    doc.add_paragraph()
    _para(doc, "Отдельная методологическая проблема — утечка данных. Varoquaux (2018) показал, "
               "что при малых выборках отбор признаков «по всем данным» до кросс-валидации "
               "завышает метрики и они не воспроизводятся. Поэтому в настоящем проекте отбор "
               "генов и масштабирование выполняются строго внутри обучающих фолдов, а "
               "«наивный» вариант приводится рядом для сравнения.")

    # 6
    _heading(doc, "6. Интерпретируемость: линейная модель и SHAP", level=2)
    _para(doc, "В медицине «чёрный ящик» неприемлем: врач должен видеть, на каких данных "
               "основан вывод. Логистическая регрессия — базовая интерпретируемая модель: "
               "вклад каждого гена равен произведению его коэффициента на стандартизованное "
               "значение. Формально эти вклады совпадают с локальными SHAP-значениями "
               "(Lundberg, Lee, 2017), что позволяет единообразно показывать и персональное "
               "объяснение (вклады генов конкретного пациента), и глобальную важность гена "
               "панели. Малая панель (15 генов) дополнительно делает вероятный клинический "
               "тест дешёвым (qPCR-формат).")

    # 7
    _heading(doc, "7. Выводы и постановка задачи", level=2)
    _bullet(doc, "терапия БА работает только на ранних стадиях — нужна доступная "
                 "диагностика стадии по крови;")
    _bullet(doc, "литература подтверждает информативность экспрессии генов крови при БА "
                 "(Lunnon 2013; Lee 2020), но работы либо бинарны (AD/контроль), либо не "
                 "дают врачу объяснения и инструмента;")
    _bullet(doc, "корректная валидация требует отбора признаков внутри фолдов (Varoquaux 2018).")
    _para(doc, "Отсюда цель работы: веб-инструмент поддержки врача, который по минимальной "
               "панели генов крови определяет стадию (норма / MCI / деменция) с честной "
               "валидацией без утечки и персональным объяснением решения.")

    # Список литературы
    doc.add_page_break()
    _heading(doc, "Список литературы", level=1)
    refs = [
        "Braak H., Braak E. (1991). Neuropathological stageing of Alzheimer-related changes. Acta Neuropathologica, 82(4), 239–259.",
        "DeTure M.A., Dickson D.W. (2019). The neuropathological diagnosis of Alzheimer's disease. Molecular Neurodegeneration, 14, 32.",
        "Folstein M.F., Folstein S.E., McHugh P.R. (1975). «Mini-mental state». Journal of Psychiatric Research, 12(3), 189–198.",
        "Hughes C.P., Berg L., Danziger W.L. et al. (1982). A new clinical scale for the staging of dementia. British Journal of Psychiatry, 140(6), 566–572.",
        "Jack C.R., Bennett D.A., Blennow K. et al. (2018). NIA-AA research framework: toward a biological definition of Alzheimer's disease. Alzheimer's & Dementia, 14(4), 535–562.",
        "Lee T., Lee H. (2020). Prediction of Alzheimer's disease using blood gene expression data. Scientific Reports, 10, 3485.",
        "Liu M., Cheng D., Wang K., Wang Y. (2018). Multi-modality cascaded convolutional neural networks for Alzheimer's disease diagnosis. Neuroinformatics, 16(3), 295–308.",
        "Lunnon K., Sattlecker M., Furney S.J. et al. (2013). A blood gene expression marker of early Alzheimer's disease. Journal of Alzheimer's Disease, 33(3), 737–753.",
        "Lundberg S.M., Lee S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765–4774.",
        "Morris J.C. (1993). The Clinical Dementia Rating (CDR): current version and scoring rules. Neurology, 43(11), 2412–2414.",
        "Petersen R.C., Doody R., Kurz A. et al. (2001). Current concepts in mild cognitive impairment. Archives of Neurology, 58(12), 1985–1992.",
        "Selkoe D.J., Hardy J. (2016). The amyloid hypothesis of Alzheimer's disease at 25 years. EMBO Molecular Medicine, 8(6), 595–608.",
        "Sims J.R., Zimmer J.A., Evans C.D. et al. (2023). Donanemab in early symptomatic Alzheimer disease: the TRAILBLAZER-ALZ 2 randomized clinical trial. JAMA, 330(6), 512–527.",
        "Spasov S., Passamonti L., Duggento A., Liò M., Toschi N. (2019). A parameter-efficient deep learning approach to predict conversion from mild cognitive impairment to Alzheimer's disease. NeuroImage, 189, 276–287.",
        "Stern Y. (2012). Cognitive reserve in ageing and Alzheimer's disease. The Lancet Neurology, 11(11), 1006–1012.",
        "van Dyck C.H., Swanson C.J., Aisen P. et al. (2023). Lecanemab in early Alzheimer's disease. New England Journal of Medicine, 388(21), 2023–2034.",
        "Varoquaux G. (2018). Cross-validation failure: small sample sizes lead to large error bars. NeuroImage, 180, 68–77.",
        "World Health Organization (2023). Dementia: key facts. WHO [Электронный ресурс].",
        "Zhang D., Wang Y., Zhou L., Yuan H., Shen D. (2011). Multimodal classification of Alzheimer's disease and mild cognitive impairment. NeuroImage, 55(3), 856–867.",
    ]
    for i, ref in enumerate(refs, 1):
        _para(doc, f"{i}. {ref}", size=11)

    path = os.path.join(DOCS, "литературный_обзор.docx")
    doc.save(path)
    print(f"  ✓ {path}")


if __name__ == "__main__":
    make_litreview()
